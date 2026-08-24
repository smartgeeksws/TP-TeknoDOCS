"""Generación del formato GCDTP-F-018 de uso de infraestructura."""

from __future__ import annotations

import re
import shutil
import tempfile
from datetime import date, datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from config.settings import INFRASTRUCTURE_TEMPLATE, OUTPUT_DIR, ROOT_DIR
from services.document_generation.confidentiality_service import DocumentGenerationError
from services.document_generation.word_pdf import PdfConversionError, convert_docx_to_pdf


class InfrastructureService:
    """Completa la plantilla institucional y produce el PDF final."""

    MONTHS = {
        1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
        5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
        9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre",
    }

    def generate(self, project: dict[str, Any], document_date: date) -> Path:
        if not INFRASTRUCTURE_TEMPLATE.is_file():
            raise DocumentGenerationError("No se encontró la plantilla institucional.")
        titular = next((talent for talent in project.get("talents", []) if talent.get("role") == "titular"), None)
        if titular is None:
            raise DocumentGenerationError("El acta requiere un talento titular asociado al proyecto.")
        expert = project.get("expert")
        if expert is None:
            raise DocumentGenerationError("El acta requiere un experto asociado al proyecto.")
        project_start_date = self._as_date(project.get("start_date"))
        if project_start_date is None:
            raise DocumentGenerationError("El proyecto no tiene fecha de inicio registrada.")

        output_directory = self._output_directory(project)
        output_directory.mkdir(parents=True, exist_ok=True)
        pdf_path = self.output_path(project)
        with tempfile.TemporaryDirectory(prefix="tp_teknodocs_") as temp_directory:
            temporary_docx = Path(temp_directory) / "uso_infraestructura.docx"
            shutil.copy2(INFRASTRUCTURE_TEMPLATE, temporary_docx)
            self._complete_template(temporary_docx, project, titular, expert, document_date, project_start_date)
            temporary_pdf = Path(temp_directory) / "uso_infraestructura.pdf"
            try:
                convert_docx_to_pdf(temporary_docx, temporary_pdf)
            except PdfConversionError as error:
                raise DocumentGenerationError(str(error)) from error
            shutil.copy2(temporary_pdf, pdf_path)
        return pdf_path

    @classmethod
    def output_path(cls, project: dict[str, Any]) -> Path:
        safe_code = cls._safe_filename(project.get("code") or "proyecto")
        return cls._output_directory(project) / f"{safe_code}_GCDTP-F-018_uso_infraestructura_compromiso.pdf"

    def _complete_template(self, document_path: Path, project: dict[str, Any], titular: dict[str, Any], expert: dict[str, Any], document_date: date, project_start_date: date) -> None:
        document = Document(document_path)
        opening = document.paragraphs[1]
        opening_text = (
            f"En la Ciudad {project['city']} a los {document_date.day} días del mes "
            f"de {self.MONTHS[document_date.month]} de {document_date.year}. Luego de "
            f"aceptado el Proyecto de Base Tecnológica: {project['code']} {project['name']}, "
            f"por el Comité de Ideas del mecanismo de intervención TecnoParque: "
            f"del {project_start_date.day} de {self.MONTHS[project_start_date.month]} "
            f"de {project_start_date.year}."
        )
        self._replace_paragraph(opening, opening_text)

        self._add_signature(document.paragraphs[31], titular.get("signature_path"))
        self._replace_paragraph(document.paragraphs[33], f"Nombre Talento: {titular['name']}")
        self._replace_paragraph(document.paragraphs[34], f"Cédula: {titular['document_type']} {titular['document_number']}")
        self._replace_paragraph(document.paragraphs[35], f"Correo electrónico: {titular.get('email') or ''}")
        self._add_signature(document.paragraphs[37], expert.get("signature_path"))
        self._replace_paragraph(document.paragraphs[39], f"Nombre Experto a cargo: {expert['name']}")
        self._replace_paragraph(document.paragraphs[40], f"Cédula: {expert['document_type']} {expert['document_number']}")
        self._replace_paragraph(document.paragraphs[41], f"Correo electrónico: {expert.get('email') or ''}")
        document.save(document_path)

    @staticmethod
    def _replace_paragraph(paragraph: Any, text: str) -> None:
        if not paragraph.runs:
            paragraph.add_run(text)
            return
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""

    @staticmethod
    def _add_signature(paragraph: Any, signature_reference: str | None) -> None:
        if not signature_reference:
            return
        signature_path = Path(signature_reference)
        if not signature_path.is_absolute():
            signature_path = ROOT_DIR / signature_path
        if not signature_path.is_file():
            raise DocumentGenerationError(f"No se encontró la firma registrada: {signature_path.name}.")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.add_run().add_picture(str(signature_path), width=Inches(1.1))

    @staticmethod
    def _safe_filename(value: str) -> str:
        cleaned = re.sub(r"[^A-Za-z0-9_-]+", "_", value.strip())
        return cleaned.strip("_") or "proyecto"

    @staticmethod
    def _as_date(value: Any) -> date | None:
        if isinstance(value, datetime):
            return value.date()
        if isinstance(value, date):
            return value
        if isinstance(value, str):
            try:
                return date.fromisoformat(value)
            except ValueError:
                return None
        return None

    @staticmethod
    def _output_directory(project: dict[str, Any]) -> Path:
        code = InfrastructureService._safe_filename(project.get("code") or str(project["id"]))
        return OUTPUT_DIR / code / "inicio"
