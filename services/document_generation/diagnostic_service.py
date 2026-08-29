"""Diligenciamiento de la plantilla Word GCDTP-F-020 V01."""

from __future__ import annotations

import io
import math
import re
import shutil
import tempfile
from copy import deepcopy
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

from config.settings import DIAGNOSTIC_TEMPLATE


class DiagnosticDocumentError(RuntimeError):
    """Error controlado al construir o validar el GCDTP-F-020."""


class DiagnosticDocumentService:
    FONT_NAME = "Calibri"
    FONT_SIZE = Pt(12)
    TABLE_FONT_SIZE = Pt(8)
    REGIONAL = "Regional Huila"
    TRAINING_CENTER = "Centro de Formación Agroindustrial La Angostura"
    TECHNOPARK = "Nodo Angostura"
    AUTO_NUMBERED_HEADINGS = {
        25,
        50,
        53,
        59,
        74,
        77,
        83,
        90,
        93,
        96,
        99,
        102,
        105,
    }
    TABLE_CAPTION_INDICES = (26, 65, 71, 80, 86)
    WORD_COUNT_PATTERN = re.compile(
        r"(?i)(?<!\w)(?:(?:\[|\()[ \t]*)?(?:[-–—][ \t]*)?"
        r"(?:[*_#]{1,3}[ \t]*){0,2}(?:la[ \t]+)?"
        r"(?:(?:(?:conteo|número|cantidad|total)[ \t]+de[ \t]+palabras)"
        r"|(?:longitud|extensión)(?:[ \t]+(?:de|del)[ \t]+"
        r"(?:texto|apartado|palabras))"
        r"|(?:(?:longitud|extensión)(?:[ \t]+(?:de|del)[ \t]+sección)?|length)(?="
        r"[ \t]*(?:[*_]{1,3}[ \t]*){0,2}:"
        r"[ \t]*(?:[*_]{1,3}[ \t]*){0,2}\d+"
        r"[ \t]*(?:[*_]{1,3}[ \t]*){0,2}(?:palabras?|words?))"
        r"|palabras|word[ \t]+count)"
        r"[ \t]*(?:[*_]{1,3}[ \t]*){0,2}:"
        r"[ \t]*(?:[*_]{1,3}[ \t]*){0,2}\d+"
        r"(?:[ \t]*(?:[*_]{1,3}[ \t]*){0,2}"
        r"(?:palabras?|words?))?"
        r"(?:[ \t]*[*_]{1,3}){0,2}[ \t]*(?:\]|\))?"
        r"(?:[ \t]*[-–—:][ \t]*)?"
    )
    HEADING_MAP = {
        25: ("1. Información general del proyecto", "Heading 1"),
        29: ("1.1 Glosario", "Heading 2"),
        32: ("1.2 Introducción", "Heading 2"),
        35: ("1.3 Planteamiento del problema o necesidad", "Heading 2"),
        38: ("1.4 Problema identificado", "Heading 2"),
        41: ("1.5 Impacto del problema", "Heading 2"),
        44: ("2. Tipo de proyecto", "Heading 1"),
        47: ("3. Objetivos del proyecto", "Heading 1"),
        50: ("3.1 Objetivo General", "Heading 2"),
        53: ("3.2 Objetivos específicos", "Heading 2"),
        56: ("4. Solución planteada", "Heading 1"),
        59: ("5. Estado del arte y de la técnica", "Heading 1"),
        62: ("5.1 Productos similares (desarrollos tecnológicos actuales)", "Heading 2"),
        68: ("5.2 Tecnología involucrada", "Heading 2"),
        74: ("6. Vigilancia tecnológica y referentes científicos", "Heading 1"),
        77: ("7. Artículos científicos relevantes y/o patentes, entre otros", "Heading 1"),
        83: ("8. Estudio legal y normativo", "Heading 1"),
        90: ("9. Análisis de viabilidad", "Heading 1"),
        93: ("10. Resultados esperados", "Heading 1"),
        96: ("11. Cronograma del proyecto", "Heading 1"),
        99: ("12. Conclusiones", "Heading 1"),
        102: ("13. Referencias bibliográficas", "Heading 2"),
        105: ("14. Anexos", "Heading 2"),
    }

    def generate(
        self,
        project: dict[str, Any],
        form_data: dict[str, Any],
        content: dict[str, Any],
    ) -> tuple[bytes, str]:
        self._validate(project, form_data, content)
        with tempfile.TemporaryDirectory(prefix="tp_diagnostic_") as directory:
            path = Path(directory) / "diagnostico.docx"
            shutil.copy2(DIAGNOSTIC_TEMPLATE, path)
            self._complete(path, project, form_data, content)
            data = path.read_bytes()
        if not data.startswith(b"PK"):
            raise DiagnosticDocumentError("El archivo Word generado no es válido.")
        code = self._safe_filename(str(project.get("code") or "proyecto"))
        return data, f"Diagnostico_Estado_del_Arte_{code}.docx"

    def _complete(
        self,
        path: Path,
        project: dict[str, Any],
        form_data: dict[str, Any],
        content: dict[str, Any],
    ) -> None:
        document = Document(path)
        if len(document.paragraphs) < 107 or len(document.tables) < 6:
            raise DiagnosticDocumentError("La estructura de la plantilla GCDTP-F-020 cambió.")

        schedule_anchor = document.paragraphs[97]
        objective_note = document.paragraphs[48]
        objective_spacing = document.paragraphs[49]
        references_paragraph = document.paragraphs[103]
        table_captions = [
            document.paragraphs[index]
            for index in self.TABLE_CAPTION_INDICES
        ]
        bibliographic_references = self._bibliographic_references(content)

        for index, (title, style) in self.HEADING_MAP.items():
            paragraph = document.paragraphs[index]
            self._replace_paragraph(
                paragraph,
                self._title_for_paragraph(paragraph, title),
                style=style,
            )

        glossary = content.get("glossary", [])
        replacements = {
            33: content["introduction"],
            36: content["problem_statement"],
            39: content["problem_identified"],
            42: content["problem_impact"],
            45: content["project_type"],
            51: content["general_objective"],
            54: "\n".join(
                f"{index}. {objective}"
                for index, objective in enumerate(content["specific_objectives"], 1)
            ),
            57: content["solution"],
            60: content["state_of_art"],
            63: "Los desarrollos comparados se seleccionaron a partir de fuentes verificadas y por su relación técnica directa con el proyecto.",
            69: content["technology_narrative"],
            75: content["technology_surveillance"],
            78: self._scientific_references_intro(
                content["scientific_references"]
            ),
            84: content["legal_study"],
            91: content["viability"],
            94: content["expected_results"],
            97: "Las actividades se organizan según la naturaleza técnica del proyecto y el periodo registrado.",
            100: content["conclusions"],
            106: "El presente documento no cuenta con anexos.",
        }
        for index, text in replacements.items():
            self._replace_paragraph(
                document.paragraphs[index],
                text,
                justify=index not in {45, 54, 78, 97, 106},
            )
        self._replace_references(
            references_paragraph,
            bibliographic_references,
        )
        self._replace_glossary(
            document.paragraphs[30],
            glossary if form_data.get("glossary_requested") else [],
        )
        for caption in table_captions:
            caption.paragraph_format.keep_with_next = True

        document.tables[0].cell(5, 0).text = "[X] Pública"
        document.tables[0].cell(5, 1).text = "[ ] Pública Clasificada"
        document.tables[0].cell(5, 2).text = "[ ] Pública Reservada"
        for cell in document.tables[0].rows[5].cells:
            self._format_cell(cell)
        self._fill_project_table(document.tables[1], project, form_data)
        self._fill_similar_products(document.tables[2], content["similar_products"])
        self._fill_technologies(document.tables[3], content["technologies"])
        self._fill_scientific_references(document.tables[4], content["scientific_references"])
        self._fill_regulations(document.tables[5], content["regulations"])
        self._insert_schedule(document, schedule_anchor, project, content["schedule"])
        for table_index, table in enumerate(document.tables):
            table_font_size = self.FONT_SIZE if table_index == 1 else self.TABLE_FONT_SIZE
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = self.FONT_NAME
                            run.font.size = table_font_size
                            run_fonts = run._element.get_or_add_rPr().rFonts
                            run_fonts.set(qn("w:eastAsia"), self.FONT_NAME)
                            run_fonts.set(qn("w:cs"), self.FONT_NAME)
        self._remove_paragraph(objective_spacing)
        self._remove_paragraph(objective_note)
        self._enable_field_updates(document)
        document.core_properties.title = "GCDTP-F-020 V01 Diagnóstico del proyecto y estado del arte"
        document.core_properties.subject = str(project.get("code") or "")
        document.save(path)
        self._validate_saved(path, content)

    def _fill_project_table(self, table: Any, project: dict[str, Any], form_data: dict[str, Any]) -> None:
        expert = project.get("expert") or {}
        team = [
            f"{item.get('name') or 'N.A.'} — {item.get('role_name') or item.get('role') or 'Talento'}"
            for item in project.get("talents", [])
        ]
        if expert.get("name"):
            team.append(f"{expert['name']} — Experto Tecnoparque")
        values = (
            project.get("name") or "N.A.",
            project.get("code") or "N.A.",
            self.TECHNOPARK,
            project.get("technology_line") or "No registrado en el proyecto",
            "\n".join(team) or "No registrado en el proyecto",
            date.fromisoformat(form_data["document_date"]).strftime("%d/%m/%Y"),
            self.TRAINING_CENTER,
            self.REGIONAL,
        )
        for row, value in enumerate(values):
            table.cell(row, 1).text = str(value)
            self._format_cell(table.cell(row, 1))

    def _fill_similar_products(self, table: Any, values: list[dict[str, Any]]) -> None:
        rows = [
            [
                item["solution"],
                f"{item['description']} {item['citation']}",
                item["relationship"],
                item["differential"],
            ]
            for item in values
        ]
        self._replace_data_rows(table, rows)

    def _fill_technologies(self, table: Any, values: list[dict[str, Any]]) -> None:
        self._replace_data_rows(table, [[item["name"], item["use"]] for item in values])

    def _fill_scientific_references(self, table: Any, values: list[dict[str, Any]]) -> None:
        self._replace_data_rows(
            table,
            [[item["type"], item["author"], str(item["year"]), f"{item['result']} {item['citation']}", item["relationship"], item["country"], item["application"]] for item in values],
        )

    def _fill_regulations(self, table: Any, values: list[dict[str, Any]]) -> None:
        self._replace_data_rows(
            table,
            [
                [
                    item["name"],
                    f"{item['application']} {item['citation']}",
                    item["intellectual_property"],
                ]
                for item in values
            ],
        )

    @classmethod
    def _scientific_references_intro(
        cls,
        values: list[dict[str, Any]],
    ) -> str:
        titles: list[str] = []
        seen: set[str] = set()
        for item in values:
            title = cls._clean_text(
                item.get("title") or item.get("result")
            ).rstrip(" .")
            key = title.casefold()
            if title and key not in seen:
                seen.add(key)
                titles.append(title)

        if not titles:
            return (
                "Se encontraron artículos científicos, patentes y/o documentos "
                "relevantes para el proyecto."
            )
        if len(titles) == 1:
            title_list = f"«{titles[0]}»"
        else:
            title_list = "; ".join(f"«{title}»" for title in titles[:-1])
            title_list += f"; y «{titles[-1]}»"
        return (
            "Se encontraron los siguientes artículos científicos, patentes y/o "
            f"documentos relevantes: {title_list}."
        )

    @classmethod
    def _bibliographic_references(
        cls,
        content: dict[str, Any],
    ) -> list[str]:
        entries: list[tuple[tuple[str, str, str], str]] = []
        seen: set[str] = set()
        for source in content.get("sources", []):
            author = cls._clean_text(source.get("author")) or "Autor no identificado"
            year = cls._clean_text(source.get("year")) or "s. f."
            title = cls._clean_text(source.get("title")) or "Fuente sin título"
            source_type = cls._clean_text(source.get("source_type"))
            reference = cls._clean_text(source.get("apa_reference"))
            reference = re.sub(
                r"(?i)\s*[.;]?\s*Recuperado(?:\s+el\s+[^,.;]+,?)?\s+de\s*:?\s*$",
                "",
                reference,
            ).rstrip(" .")
            if not reference:
                reference = f"{author} ({year}). {title}"
                if source_type:
                    reference += f". {source_type}"

            url = cls._clean_text(source.get("url"))
            if url:
                reference = f"{reference}. Recuperado de {url}"

            identity = url.casefold() or re.sub(
                r"\W+", "", reference
            ).casefold()
            if identity in seen:
                continue
            seen.add(identity)
            entries.append(
                ((author.casefold(), year.casefold(), title.casefold()), reference)
            )

        if not entries:
            for reference_value in content.get("references", []):
                reference = cls._clean_text(reference_value)
                identity = re.sub(r"\W+", "", reference).casefold()
                if reference and identity not in seen:
                    seen.add(identity)
                    entries.append(((identity, "", ""), reference))

        references = [
            reference
            for _sort_key, reference in sorted(entries, key=lambda item: item[0])
        ]
        if not references:
            raise DiagnosticDocumentError(
                "No se encontraron fuentes para construir las referencias bibliográficas."
            )
        return references

    def _insert_schedule(
        self,
        document: Document,
        anchor: Any,
        project: dict[str, Any],
        schedule: list[dict[str, Any]],
    ) -> None:
        table = document.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        self._set_table_cell_margins(table)
        self._prevent_row_split(table.rows[0])
        self._repeat_table_header(table.rows[0])
        headers = (
            "Fase",
            "Actividad",
            "Fecha inicial",
            "Fecha final",
            "Duración",
            "Entregable",
        )
        for index, header in enumerate(headers):
            table.cell(0, index).text = header
            self._format_cell(table.cell(0, index))
        dates = self._schedule_dates(project, len(schedule))
        for item, date_range in zip(schedule, dates):
            row = table.add_row()
            self._prevent_row_split(row)
            cells = row.cells
            start, end = date_range
            values = (
                item["phase"],
                item["activity"],
                start.strftime("%d/%m/%Y") if start else "Por definir",
                end.strftime("%d/%m/%Y") if end else "Por definir",
                str((end - start).days + 1) if start and end else "Por definir",
                item["deliverable"],
            )
            for cell, value in zip(cells, values):
                cell.text = self._clean_text(value)
                self._format_cell(cell)
        self._set_schedule_widths(document, table)
        anchor._p.addnext(table._tbl)

    @staticmethod
    def _set_table_cell_margins(table: Any) -> None:
        properties = table._tbl.tblPr
        margins = properties.find(qn("w:tblCellMar"))
        if margins is None:
            margins = OxmlElement("w:tblCellMar")
            properties.append(margins)
        for name, value in (
            ("top", 20),
            ("left", 40),
            ("bottom", 20),
            ("right", 40),
        ):
            margin = margins.find(qn(f"w:{name}"))
            if margin is None:
                margin = OxmlElement(f"w:{name}")
                margins.append(margin)
            margin.set(qn("w:w"), str(value))
            margin.set(qn("w:type"), "dxa")

    @staticmethod
    def _set_schedule_widths(document: Document, table: Any) -> None:
        section = document.sections[0]
        available_width = int(
            section.page_width - section.left_margin - section.right_margin
        )
        ratios = (0.175, 0.21, 0.15, 0.15, 0.12, 0.195)
        widths = [int(available_width * ratio) for ratio in ratios[:-1]]
        widths.append(available_width - sum(widths))
        table.autofit = False
        for index, width in enumerate(widths):
            table.columns[index].width = width
            for cell in table.columns[index].cells:
                cell.width = width

    @staticmethod
    def _schedule_dates(
        project: dict[str, Any],
        count: int,
    ) -> list[tuple[date | None, date | None]]:
        start = DiagnosticDocumentService._as_date(project.get("start_date"))
        end = DiagnosticDocumentService._as_date(project.get("end_date"))
        if start is None or end is None or end < start:
            return [(None, None)] * count
        total = (end - start).days + 1
        ranges = []
        for index in range(count):
            start_offset = math.floor(index * total / count)
            end_offset = max(start_offset, math.floor((index + 1) * total / count) - 1)
            ranges.append((start + timedelta(days=start_offset), min(end, start + timedelta(days=end_offset))))
        return ranges

    @staticmethod
    def _replace_data_rows(table: Any, rows: list[list[str]]) -> None:
        prototype = deepcopy(table.rows[1]._tr)
        for row in list(table.rows)[1:]:
            table._tbl.remove(row._tr)
        for values in rows:
            table._tbl.append(deepcopy(prototype))
            current = table.rows[-1]
            DiagnosticDocumentService._prevent_row_split(current)
            for cell, value in zip(current.cells, values):
                cell.text = DiagnosticDocumentService._clean_text(value)
                DiagnosticDocumentService._format_cell(cell)

    @staticmethod
    def _repeat_table_header(row: Any) -> None:
        properties = row._tr.get_or_add_trPr()
        if properties.find(qn("w:tblHeader")) is None:
            properties.append(OxmlElement("w:tblHeader"))

    @staticmethod
    def _prevent_row_split(row: Any) -> None:
        properties = row._tr.get_or_add_trPr()
        if properties.find(qn("w:cantSplit")) is None:
            properties.append(OxmlElement("w:cantSplit"))

    @staticmethod
    def _format_cell(cell: Any) -> None:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                DiagnosticDocumentService._format_run(run)

    @staticmethod
    def _replace_paragraph(paragraph: Any, text: str, *, style: str | None = None, justify: bool = False) -> None:
        original_run_properties = None
        if style:
            original_run = next(
                (run for run in paragraph.runs if run.text.strip()),
                None,
            )
            if original_run is not None and original_run._r.rPr is not None:
                original_run_properties = deepcopy(original_run._r.rPr)

        for child in list(paragraph._p):
            if child.tag != qn("w:pPr"):
                paragraph._p.remove(child)
        run = paragraph.add_run(DiagnosticDocumentService._clean_text(text))
        if original_run_properties is not None:
            run._r.insert(0, original_run_properties)
        if style:
            paragraph.style = style
            paragraph.paragraph_format.keep_with_next = True
        else:
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.JUSTIFY
                if justify
                else WD_ALIGN_PARAGRAPH.LEFT
            )
        DiagnosticDocumentService._format_run(run)

    @classmethod
    def _replace_glossary(cls, paragraph: Any, glossary: list[dict[str, Any]]) -> None:
        for child in list(paragraph._p):
            if child.tag != qn("w:pPr"):
                paragraph._p.remove(child)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(6)
        if not glossary:
            run = paragraph.add_run("No aplica para el presente proyecto.")
            cls._format_run(run)
            return

        paragraph_properties = (
            deepcopy(paragraph._p.pPr)
            if paragraph._p.pPr is not None
            else None
        )
        current = paragraph
        for index, item in enumerate(glossary):
            if index:
                element = OxmlElement("w:p")
                if paragraph_properties is not None:
                    element.append(deepcopy(paragraph_properties))
                current._p.addnext(element)
                current = Paragraph(element, paragraph._parent)
            current.alignment = WD_ALIGN_PARAGRAPH.LEFT
            current.paragraph_format.keep_together = True
            current.paragraph_format.space_after = Pt(6)
            term = current.add_run(f"{cls._clean_text(item['term'])}: ")
            term.bold = True
            cls._format_run(term)
            definition = cls._clean_text(
                f"{item['definition']} {item.get('citation', '')}"
            )
            run = current.add_run(definition)
            cls._format_run(run)

    @classmethod
    def _replace_references(
        cls,
        paragraph: Any,
        references: list[str],
    ) -> None:
        paragraph_properties = (
            deepcopy(paragraph._p.pPr)
            if paragraph._p.pPr is not None
            else None
        )
        for child in list(paragraph._p):
            if child.tag != qn("w:pPr"):
                paragraph._p.remove(child)

        current = paragraph
        for index, reference in enumerate(references):
            if index:
                element = OxmlElement("w:p")
                if paragraph_properties is not None:
                    element.append(deepcopy(paragraph_properties))
                current._p.addnext(element)
                current = Paragraph(element, paragraph._parent)
            current.alignment = WD_ALIGN_PARAGRAPH.LEFT
            current.paragraph_format.left_indent = Inches(0.5)
            current.paragraph_format.first_line_indent = Inches(-0.5)
            current.paragraph_format.line_spacing = 2
            current.paragraph_format.space_after = Pt(0)
            current.paragraph_format.keep_together = True
            run = current.add_run(cls._clean_text(reference))
            cls._format_run(run)

    @classmethod
    def _format_run(cls, run: Any) -> None:
        run.font.name = cls.FONT_NAME
        run.font.size = cls.FONT_SIZE
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), cls.FONT_NAME)

    @classmethod
    def _clean_text(cls, value: Any) -> str:
        marker = "\ue000"
        text = cls.WORD_COUNT_PATTERN.sub(marker, str(value or ""))
        text = re.sub(
            rf"(?m)^[^\w\r\n]*{marker}"
            rf"(?:[^\w\r\n]*{marker})*[^\w\r\n]*(?:\r?\n|$)",
            "",
            text,
        )
        text = text.replace(marker, "")
        text = re.sub(r"[ \t]{2,}", " ", text)
        text = re.sub(r"[ \t]+([,.;:!?])", r"\1", text)
        text = re.sub(r"([.!?])\1+", r"\1", text)
        text = re.sub(r"([,;:])\1+", r"\1", text)
        text = re.sub(r"[ \t]+\n", "\n", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        if text and not re.search(r"\w", text):
            return ""
        return text.strip()

    @staticmethod
    def _title_for_paragraph(paragraph: Any, title: str) -> str:
        num_pr = paragraph._p.pPr.numPr if paragraph._p.pPr is not None else None
        if num_pr is None:
            return title
        return re.sub(r"^\d+(?:\.\d+)*\.?\s+", "", title).strip()

    @staticmethod
    def _remove_paragraph(paragraph: Any) -> None:
        element = paragraph._element
        element.getparent().remove(element)
        paragraph._p = paragraph._element = None

    @staticmethod
    def _enable_field_updates(document: Document) -> None:
        settings = document.settings._element
        update = settings.find(qn("w:updateFields"))
        if update is None:
            update = OxmlElement("w:updateFields")
            settings.append(update)
        update.set(qn("w:val"), "true")
        for field in document.element.body.xpath('.//w:fldChar[@w:fldCharType="begin"]'):
            field.set(qn("w:dirty"), "true")

    def _validate_saved(self, path: Path, content: dict[str, Any]) -> None:
        document = Document(path)
        full_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        paragraph_texts = {paragraph.text.strip() for paragraph in document.paragraphs}
        missing = []
        for index, (title, _) in self.HEADING_MAP.items():
            expected = (
                re.sub(r"^\d+(?:\.\d+)*\.?\s+", "", title).strip()
                if index in self.AUTO_NUMBERED_HEADINGS
                else title
            )
            if expected not in paragraph_texts:
                missing.append(title)
        if missing:
            raise DiagnosticDocumentError("Faltan títulos en el documento generado.")
        if self.WORD_COUNT_PATTERN.search(full_text):
            raise DiagnosticDocumentError("El documento conserva un contador interno de palabras.")
        if "(si aplica)" in full_text.lower():
            raise DiagnosticDocumentError("El glosario conserva la indicación '(si aplica)'.")
        if "El presente documento no cuenta con anexos." not in full_text:
            raise DiagnosticDocumentError("No se diligenció correctamente el apartado de anexos.")
        if len(document.tables) < 7:
            raise DiagnosticDocumentError("No se generó la tabla de cronograma.")
        scientific_intro = self._scientific_references_intro(
            content["scientific_references"]
        )
        if scientific_intro not in paragraph_texts:
            raise DiagnosticDocumentError(
                "No se incluyeron los títulos de los referentes científicos."
            )
        missing_references = [
            reference
            for reference in self._bibliographic_references(content)
            if reference not in paragraph_texts
        ]
        if missing_references:
            raise DiagnosticDocumentError(
                "No se incluyeron todas las referencias bibliográficas."
            )

    @staticmethod
    def _validate(project: dict[str, Any], form_data: dict[str, Any], content: dict[str, Any]) -> None:
        if not DIAGNOSTIC_TEMPLATE.is_file():
            raise DiagnosticDocumentError("No se encontró la plantilla GCDTP-F-020.")
        for key in ("id", "code", "name", "description"):
            if not project.get(key):
                raise DiagnosticDocumentError(f"Falta el dato del proyecto: {key}.")
        required_content = {
            "introduction", "problem_statement", "problem_identified",
            "problem_impact", "project_type", "general_objective",
            "specific_objectives", "solution", "state_of_art",
            "similar_products", "technology_narrative", "technologies",
            "technology_surveillance", "scientific_references", "legal_study",
            "regulations", "viability", "expected_results", "schedule",
            "conclusions", "references", "sources",
        }
        missing = sorted(required_content - set(content))
        if missing:
            raise DiagnosticDocumentError(
                "Falta contenido para diligenciar: " + ", ".join(missing) + "."
            )
    @staticmethod
    def _as_date(value: Any) -> date | None:
        if isinstance(value, datetime):
            return value.date()
        if isinstance(value, date):
            return value
        if isinstance(value, str) and value:
            try:
                return date.fromisoformat(value)
            except ValueError:
                return None
        return None

    @staticmethod
    def _safe_filename(value: str) -> str:
        cleaned = re.sub(r"[^A-Za-z0-9_-]+", "_", value.strip())
        return cleaned.strip("_") or "proyecto"
