"""Generación del Acuerdo de Confidencialidad y Compromisos GCDTP-F-017."""

from __future__ import annotations

import re
import shutil
import tempfile
from copy import deepcopy
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from config.settings import (
    CONFIDENTIALITY_TEMPLATE,
    FIXED_SIGNATURES_DIR,
    OUTPUT_DIR,
)
from services.document_generation.word_pdf import (
    PdfConversionError,
    convert_docx_to_pdf,
)
from services.signature_storage import SignatureStorage, SignatureStorageError

MONTHS = {
    1: "enero",
    2: "febrero",
    3: "marzo",
    4: "abril",
    5: "mayo",
    6: "junio",
    7: "julio",
    8: "agosto",
    9: "septiembre",
    10: "octubre",
    11: "noviembre",
    12: "diciembre",
}
ROLE_LABELS = {
    "titular": "Titular",
    "interlocutor": "Interlocutor",
    "ejecutor": "Ejecutor",
}
FIXED_SIGNERS = [
    {
        "name": "Sergio Andrés Cabrera",
        "role": "Nombre del Experto Tecnoparque",
        "document": "C.C. 1.110.454.504",
        "signature": "fsergio.png",
    },
    {
        "name": "Carolina Garcia Monje",
        "role": "Nombre del Experto Tecnoparque",
        "document": "C.C. 36.301.495",
        "signature": "fcaro.png",
    },
    {
        "name": "Diego Alfonso Polania",
        "role": "Nombre del Experto Tecnoparque",
        "document": "C.C. 7.684.683",
        "signature": "fdiego.png",
    },
    {
        "name": "Cesar Augusto Pérez Tafur",
        "role": "Nombre del Experto Tecnoparque",
        "document": "C.C. 7.728.013",
        "signature": "fcesar.png",
    },
    {
        "name": "Maria Andrea Qimbaya",
        "role": "Nombre del Apoyo Administrativo",
        "document": "C.C. 1003.812.026",
        "signature": "fmaria.png",
    },
    {
        "name": "Lina Marcela Trujillo Osso",
        "role": "SUBDIRECTORA DE CENTRO G02(E)",
        "document": "C.C. 52.701.590",
        "signature": None,
    },
]


class DocumentGenerationError(PdfConversionError):
    """Error controlado durante la generación documental."""


class ConfidentialityService:
    """Completa la plantilla institucional y produce el PDF final."""

    def generate(self, project: dict[str, Any], document_date: date) -> Path:
        if not CONFIDENTIALITY_TEMPLATE.is_file():
            raise DocumentGenerationError("No se encontró la plantilla institucional.")

        signers = self._talent_signers(project)
        output_directory = self._output_directory(project)
        output_directory.mkdir(parents=True, exist_ok=True)
        pdf_path = self.output_path(project)

        with tempfile.TemporaryDirectory(prefix="tp_teknodocs_") as temp_directory:
            temporary_docx = Path(temp_directory) / "confidencialidad.docx"
            shutil.copy2(CONFIDENTIALITY_TEMPLATE, temporary_docx)
            self._complete_template(
                temporary_docx,
                project,
                document_date,
                signers,
            )
            temporary_pdf = Path(temp_directory) / "confidencialidad.pdf"
            try:
                convert_docx_to_pdf(temporary_docx, temporary_pdf)
            except PdfConversionError as error:
                raise DocumentGenerationError(str(error)) from error
            shutil.copy2(temporary_pdf, pdf_path)

        return pdf_path

    @classmethod
    def output_path(cls, project: dict[str, Any]) -> Path:
        """Ruta estable utilizada también para determinar el estado documental."""

        safe_code = cls._safe_filename(project.get("code") or "proyecto")
        return cls._output_directory(project) / (
            f"{safe_code}_GCDTP-F-017_confidencialidad_compromiso.pdf"
        )

    def _complete_template(
        self,
        document_path: Path,
        project: dict[str, Any],
        document_date: date,
        talent_signers: list[dict[str, Any]],
    ) -> None:
        document = Document(document_path)
        self._replace_paragraph(
            document.paragraphs[1],
            self._opening_text(project, document_date, talent_signers),
        )
        closing = (
            f"Para constancia, se firma el documento en la ciudad de "
            f"{project['city']} a los {document_date.day} días del mes de "
            f"{MONTHS[document_date.month]} de {document_date.year}, por las partes:"
        )
        closing_paragraph = next(
            (
                paragraph
                for paragraph in document.paragraphs
                if paragraph.text.startswith("Para constancia, se firma")
            ),
            None,
        )
        if closing_paragraph is None:
            raise DocumentGenerationError(
                "La plantilla no contiene el párrafo final esperado."
            )
        self._replace_paragraph(closing_paragraph, closing)
        self._replace_signature_table(document, talent_signers)
        document.save(document_path)

    def _talent_signers(self, project: dict[str, Any]) -> list[dict[str, Any]]:
        talents_by_role = {
            talent["role"]: talent
            for talent in project.get("talents", [])
        }
        titular = talents_by_role.get("titular")
        if titular is None:
            raise DocumentGenerationError(
                "El acta requiere un talento titular asociado al proyecto."
            )

        effective_roles = dict(talents_by_role)
        if len(effective_roles) == 1:
            effective_roles["interlocutor"] = titular
            effective_roles["ejecutor"] = titular
        elif "interlocutor" not in effective_roles:
            effective_roles["interlocutor"] = titular
        elif "ejecutor" not in effective_roles:
            effective_roles["ejecutor"] = titular

        grouped: dict[int, dict[str, Any]] = {}
        for role in ("titular", "interlocutor", "ejecutor"):
            talent = effective_roles[role]
            signer = grouped.setdefault(
                talent["id"],
                {"person": talent, "roles": []},
            )
            signer["roles"].append(role)

        return list(grouped.values())

    def _opening_text(
        self,
        project: dict[str, Any],
        document_date: date,
        talent_signers: list[dict[str, Any]],
    ) -> str:
        participants = []
        for signer in talent_signers:
            person = signer["person"]
            roles = self._join_roles(signer["roles"])
            participants.append(
                f"{person['name']}, identificado(a) con "
                f"{person['document_type']} N° {person['document_number']} de "
                f"{person['document_issue_place']}, quien actúa como talento {roles}"
            )
        participant_text = "; y ".join(participants)
        return (
            f"En la ciudad de {project['city']} a los {document_date.day} días del mes "
            f"de {MONTHS[document_date.month]} de {document_date.year}, se celebra el "
            f"presente Acuerdo de Confidencialidad y Compromisos entre el SENA, "
            f"representado por los firmantes abajo de este documento, y por otra parte "
            f"{participant_text}, del proyecto «{project['name']}» con número de "
            f"consecutivo de la idea aprobada por comité Número: {project['code']}, "
            f"previas las siguientes consideraciones:"
        )

    def _replace_signature_table(
        self,
        document: Document,
        talent_signers: list[dict[str, Any]],
    ) -> None:
        if not document.tables:
            raise DocumentGenerationError("La plantilla no contiene la tabla de firmas.")
        table = document.tables[0]
        template_row = deepcopy(table.rows[5]._tr)
        for row in list(table.rows)[1:]:
            table._tbl.remove(row._tr)

        signers = []
        for talent in talent_signers:
            person = talent["person"]
            signature_path = self._resolve_signature(person.get("signature_path"))
            signers.append(
                {
                    "name": person["name"],
                    "role": f"Talento {self._join_roles(talent['roles'])}",
                    "document": (
                        f"{person['document_type']} {person['document_number']}"
                    ),
                    "signature_path": signature_path,
                }
            )
        for fixed in FIXED_SIGNERS:
            signature_path = (
                FIXED_SIGNATURES_DIR / fixed["signature"]
                if fixed["signature"]
                else None
            )
            if signature_path and not signature_path.is_file():
                raise DocumentGenerationError(
                    f"No se encontró la firma institucional {fixed['signature']}."
                )
            signers.append({**fixed, "signature_path": signature_path})

        for signer in signers:
            new_row = deepcopy(template_row)
            table._tbl.append(new_row)
            self._fill_signature_row(table.rows[-1], signer)

    def _fill_signature_row(self, row: Any, signer: dict[str, Any]) -> None:
        row.height = None
        cell = row.cells[0]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = ""
        nested = cell.add_table(rows=1, cols=2)
        nested.autofit = False
        nested.columns[0].width = Inches(4.0)
        nested.columns[1].width = Inches(2.0)
        self._remove_table_borders(nested)

        info_cell, signature_cell = nested.rows[0].cells
        info = info_cell.paragraphs[0]
        info.paragraph_format.space_after = Pt(0)
        name_run = info.add_run(signer["name"])
        name_run.bold = True
        name_run.font.size = Pt(12)
        info.add_run(f"\n{signer['role']}").font.size = Pt(12)
        info.add_run(f"\n{signer['document']}").font.size = Pt(12)

        signature_paragraph = signature_cell.paragraphs[0]
        signature_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        signature_path = signer.get("signature_path")
        if signature_path:
            image = (
                str(signature_path)
                if isinstance(signature_path, Path)
                else signature_path
            )
            signature_paragraph.add_run().add_picture(
                image,
                width=Inches(1.65),
            )
        label = signature_cell.add_paragraph("Firma")
        label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label.runs[0].bold = True
        label.runs[0].font.size = Pt(12)

    @staticmethod
    def _remove_table_borders(table: Any) -> None:
        properties = table._tbl.tblPr
        borders = properties.first_child_found_in("w:tblBorders")
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            properties.append(borders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            element = OxmlElement(f"w:{edge}")
            element.set(qn("w:val"), "nil")
            borders.append(element)

    @staticmethod
    def _replace_paragraph(paragraph: Any, text: str) -> None:
        if not paragraph.runs:
            paragraph.add_run(text)
            return
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""

    @staticmethod
    def _join_roles(roles: list[str]) -> str:
        labels = [ROLE_LABELS[role] for role in roles]
        if len(labels) == 1:
            return labels[0]
        if len(labels) == 2:
            return f"{labels[0]} e {labels[1]}"
        return f"{labels[0]}, {labels[1]} y {labels[2]}"

    @staticmethod
    def _resolve_signature(signature_reference: str | None) -> Any:
        if not signature_reference:
            raise DocumentGenerationError(
                "Uno de los talentos no tiene firma registrada."
            )
        try:
            return SignatureStorage().open(signature_reference)
        except SignatureStorageError as error:
            raise DocumentGenerationError(str(error)) from error

    @staticmethod
    def _safe_filename(value: str) -> str:
        cleaned = re.sub(r"[^A-Za-z0-9_-]+", "_", value.strip())
        return cleaned.strip("_") or "proyecto"

    @staticmethod
    def _output_directory(project: dict[str, Any]) -> Path:
        code = ConfidentialityService._safe_filename(
            project.get("code") or str(project["id"])
        )
        return OUTPUT_DIR / code / "inicio"
