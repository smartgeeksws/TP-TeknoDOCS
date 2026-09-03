"""Document generation for the closure phase."""

from __future__ import annotations

from datetime import date
from html import escape
from io import BytesIO
from pathlib import Path
import re
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt

from config.settings import FINAL_REPORT_TEMPLATE


class ClosureDocumentError(RuntimeError):
    """Controlled error while creating a closure document."""


def safe_filename(value: str) -> str:
    return re.sub(r'[<>:"/\\|?*]', "-", str(value)).strip() or "proyecto"


def format_date(value: Any) -> str:
    if not value:
        return "No registrada"
    if isinstance(value, date):
        return value.strftime("%d/%m/%Y")
    try:
        return date.fromisoformat(str(value)).strftime("%d/%m/%Y")
    except ValueError:
        return str(value)


class FinalReportDocumentService:
    """Fills the official GCDTP-F-023 template without modifying it."""

    FIELD_HEADINGS = {
        "Introducción": "introduccion",
        "Planteamiento del problema": "planteamiento_problema",
        "4.1 Objetivo General": "objetivo_general",
        "4.2 Objetivos Específicos": "objetivos_especificos",
        "5. Estado del arte y estado de la técnica": "estado_arte",
        "6. Metodología de desarrollo": "metodologia",
        "7. Desarrollo del proyecto": "desarrollo",
        "8. Resultados obtenidos": "resultados",
        "9. Análisis de viabilidad": "analisis_viabilidad",
        "10. Propiedad intelectual y transferencia tecnológica": "propiedad_transferencia",
        "11. Impacto del proyecto": "impacto",
        "12. Conclusiones": "conclusiones",
        "13. Referencias bibliográficas": "referencias",
        "14. Anexos": "anexos",
    }
    RENAMED_HEADINGS = {
        "8. Resultados obtenidos": "9. Resultados obtenidos",
        "9. Análisis de viabilidad": "10. Análisis de viabilidad",
        "10. Propiedad intelectual y transferencia tecnológica": (
            "11. Propiedad intelectual y transferencia tecnológica"
        ),
        "11. Impacto del proyecto": "12. Impacto del proyecto",
        "12. Conclusiones": "13. Conclusiones",
        "13. Referencias bibliográficas": "14. Referencias bibliográficas",
        "14. Anexos": "15. Anexos",
    }

    def generate(
        self, project: dict[str, Any], form_data: dict[str, Any], content: dict[str, str]
    ) -> tuple[bytes, str]:
        if not Path(FINAL_REPORT_TEMPLATE).is_file():
            raise ClosureDocumentError("No se encontro la plantilla GCDTP-F-023.")
        document = Document(FINAL_REPORT_TEMPLATE)
        self._fill_identification(document, project, form_data)
        self._clear_objectives_intro(document)
        for heading, field in self.FIELD_HEADINGS.items():
            self._replace_after_heading(document, heading, content.get(field, ""))
        self._insert_normativity(document, content.get("normatividad", ""))
        self._renumber_headings(document)
        self._remove_instructions(document)
        output = BytesIO()
        document.save(output)
        return output.getvalue(), f"Informe Final {safe_filename(project.get('code', 'proyecto'))}.docx"

    def _fill_identification(self, document: Document, project: dict[str, Any], data: dict[str, Any]) -> None:
        if len(document.tables) < 2:
            raise ClosureDocumentError("La plantilla GCDTP-F-023 no tiene la estructura esperada.")
        tables = document.tables
        classification = tables[0].rows[5].cells[0]
        if "X" not in classification.text:
            classification.paragraphs[0].add_run(" X")
        talent = next(
            (item for item in project.get("talents", []) if item.get("role") == "titular"),
            None,
        ) or {}
        values = [
            talent.get("name") or "No registrado",
            project.get("name") or "No registrado",
            project.get("code") or "No registrado",
            (project.get("expert") or {}).get("name") or "No registrado",
            project.get("technology_line") or "No registrada",
            project.get("initial_trl") or "No registrado",
            data.get("achieved_trl") or project.get("target_trl") or "No registrado",
            project.get("city") or "No registrada",
            format_date(data.get("delivery_date")),
        ]
        for row, value in zip(tables[1].rows, values):
            self._set_cell(row.cells[1], str(value))

    @staticmethod
    def _set_cell(cell: Any, value: str) -> None:
        paragraph = cell.paragraphs[0]
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = value
        else:
            paragraph.add_run(value)

    def _replace_after_heading(self, document: Document, heading: str, value: str) -> None:
        normalized_heading = self._normalize(heading)
        paragraphs = document.paragraphs
        for index, paragraph in enumerate(paragraphs[:-1]):
            if self._normalize(paragraph.text) == normalized_heading:
                next_paragraph = paragraphs[index + 1]
                if heading in {"4.1 Objetivo General", "4.2 Objetivos Específicos"}:
                    replacement = document.add_paragraph()
                    paragraph._p.addnext(replacement._p)
                    self._replace_paragraph(replacement, value)
                else:
                    self._replace_paragraph(next_paragraph, value)
                return

    def _clear_objectives_intro(self, document: Document) -> None:
        self._replace_after_heading(document, "Objetivos", "")

    def _insert_normativity(self, document: Document, value: str) -> None:
        paragraphs = document.paragraphs
        for index, paragraph in enumerate(paragraphs[:-1]):
            if self._normalize(paragraph.text) == self._normalize("7. Desarrollo del proyecto"):
                heading = document.add_paragraph()
                paragraphs[index + 1]._p.addnext(heading._p)
                heading.add_run("8. Normatividad").bold = True
                body = document.add_paragraph()
                heading._p.addnext(body._p)
                self._replace_paragraph(body, value)
                return

    def _renumber_headings(self, document: Document) -> None:
        for paragraph in document.paragraphs:
            for old_heading, new_heading in self.RENAMED_HEADINGS.items():
                if self._normalize(paragraph.text) == self._normalize(old_heading):
                    self._replace_paragraph(paragraph, new_heading)
                    break

    @staticmethod
    def _replace_paragraph(paragraph: Any, value: str) -> None:
        for run in paragraph.runs:
            run.text = ""
        lines = value.splitlines() or [""]
        paragraph.add_run(lines[0])
        for line in lines[1:]:
            paragraph.add_run().add_break()
            paragraph.add_run(line)

    def _remove_instructions(self, document: Document) -> None:
        for paragraph in list(document.paragraphs):
            if self._normalize(paragraph.text) == "instrucciones":
                element = paragraph._element
                parent = element.getparent()
                remove = False
                for sibling in list(parent):
                    if sibling is element:
                        remove = True
                    if remove and sibling.tag.endswith("}p"):
                        parent.remove(sibling)
                return

    @staticmethod
    def _normalize(value: str) -> str:
        replacements = str.maketrans("áéíóúüñ", "aeiouun")
        return " ".join(value.lower().translate(replacements).split())


class CertificationLetterDocumentService:
    """Creates the requested plain formal certification letter."""

    TITLE = "CARTA DE CERTIFICACIÓN Y PARTICIPACIÓN EN CONSULTORÍA CIENTÍFICO-TECNOLÓGICA"

    def generate(self, values: dict[str, str], technical: dict[str, str]) -> tuple[bytes, str]:
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        style = document.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(self.TITLE)
        title_run.bold = True
        title_run.font.name = "Arial"
        title_run.font.size = Pt(11)
        self._paragraph(document, f"{values['city']}, {format_date(values['issue_date'])}")
        self._paragraph(
            document,
            "Yo, {beneficiary_name}, identificado(a) con {document_type} No. "
            "{document_number}, en calidad de beneficiario(a) del acompañamiento "
            "recibido, certifico que participé en el proceso de consultoría "
            "científico-tecnológica brindado por TecnoParque Nodo Angostura, a "
            "través del experto {expert_name}, quien desempeñó el rol de {expert_role}.".format(**values),
        )
        self._paragraph(document, "El acompañamiento se realizó en el marco del proyecto denominado:")
        named = document.add_paragraph()
        named.alignment = WD_ALIGN_PARAGRAPH.CENTER
        named.add_run(f'“{values["project_name"]}”, identificado con el código {values["project_code"]}.').bold = True
        self._paragraph(document, "El proceso de acompañamiento inició el día {start_date} y finalizó el día {end_date}, con el objetivo de {objective}.".format(objective=technical["objetivo_acompanamiento"], **values))
        self._paragraph(document, "Durante la consultoría se brindó asesoría y acompañamiento técnico en " + technical["descripcion_consultoria"] + ".")
        self._paragraph(document, "Como resultado del acompañamiento, el proyecto alcanzó un nivel de madurez tecnológica {achieved_trl}, {result}.".format(result=technical["resultado_trl"], **values))
        self._paragraph(document, "Manifiesto que la calidad del acompañamiento recibido fue {rating}, destacando la pertinencia técnica, la claridad en las orientaciones, el cumplimiento de los objetivos propuestos y el aporte realizado al fortalecimiento del proyecto.".format(**values))
        self._paragraph(document, "La presente certificación se expide a solicitud de la persona interesada, para los fines que estime pertinentes.")
        self._paragraph(document, "Atentamente,")
        self._paragraph(document, "\n{beneficiary_name}\n{document_type} No. {document_number}\nBeneficiario(a) del proyecto".format(**values))
        if values.get("company"):
            self._paragraph(document, values["company"])
        self._paragraph(document, values["city"])
        output = BytesIO()
        document.save(output)
        return output.getvalue(), f"Carta_Certificacion_{safe_filename(values['project_code'])}.docx"

    @staticmethod
    def _paragraph(document: Document, text: str) -> None:
        paragraph = document.add_paragraph(text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(10)
        paragraph.paragraph_format.line_spacing = 1.15


class BusinessModelPdfService:
    """Produces the landscape PDF layout used for the business model."""

    LABELS = {
        "propuesta_valor": "Propuesta de valor",
        "segmento_clientes_adopcion": "Segmento de clientes y estrategia de adopción",
        "canales_distribucion": "Canales de distribución",
        "relaciones_clientes": "Relaciones con clientes",
        "flujo_ingresos": "Flujo de ingresos",
        "recursos_clave": "Recursos clave",
        "actividades_clave": "Actividades clave",
        "alianzas_clave": "Alianzas clave",
        "estructura_costos": "Estructura de costos",
    }

    def generate(self, project: dict[str, Any], content: dict[str, str]) -> tuple[bytes, str]:
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import landscape, letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
        except ImportError as error:
            raise ClosureDocumentError("La dependencia reportlab no esta instalada.") from error
        output = BytesIO()
        document = SimpleDocTemplate(output, pagesize=landscape(letter), rightMargin=0.35 * inch, leftMargin=0.35 * inch, topMargin=0.35 * inch, bottomMargin=0.35 * inch)
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("closure-title", parent=styles["Title"], fontSize=16, leading=19, textColor=colors.HexColor("#00304D"), alignment=1)
        section_style = ParagraphStyle("closure-section", parent=styles["Heading2"], fontSize=10, leading=12, textColor=colors.HexColor("#007832"), spaceBefore=6)
        body_style = ParagraphStyle("closure-body", parent=styles["BodyText"], fontSize=8, leading=10)
        canvas_base_style = ParagraphStyle("closure-canvas", parent=body_style)
        story = [Paragraph("Modelo de Negocio - Lean Canvas", title_style), Spacer(1, 8), Paragraph(f"<b>Proyecto:</b> {escape(str(project.get('name') or 'No registrado'))}", body_style), Spacer(1, 4)]
        for field, label in self.LABELS.items():
            story.extend([Paragraph(label, section_style), Paragraph(escape(content[field]).replace("\n", "<br/>"), body_style)])
        cells = []
        order = list(self.LABELS)
        for field in order:
            word_count = max(len(content[field].split()), 1)
            # Shorter blocks use a larger type size while long blocks remain legible.
            font_size = max(5.5, min(7.2, 1050 / word_count))
            canvas_style = ParagraphStyle(
                f"closure-canvas-{field}",
                parent=canvas_base_style,
                fontSize=font_size,
                leading=font_size + 1.1,
            )
            cells.append(Paragraph(f"<b>{self.LABELS[field]}</b><br/>{escape(content[field]).replace(chr(10), '<br/>')}", canvas_style))
        rows = [cells[0:3], cells[3:6], cells[6:9]]
        table = Table(rows, colWidths=[3.25 * inch] * 3, rowHeights=[2.25 * inch] * 3)
        table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#00304D")), ("VALIGN", (0, 0), (-1, -1), "TOP"), ("BACKGROUND", (0, 0), (-1, -1), colors.white), ("LEFTPADDING", (0, 0), (-1, -1), 5), ("RIGHTPADDING", (0, 0), (-1, -1), 5), ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 4)]))
        story.extend([PageBreak(), Paragraph("Lean Canvas", title_style), Spacer(1, 6), table])
        document.build(story)
        return output.getvalue(), f"Modelo_Negocios_{safe_filename(project.get('code', 'proyecto'))}.pdf"
