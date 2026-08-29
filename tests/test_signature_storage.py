from __future__ import annotations

import base64
import shutil
import tempfile
import unittest
from datetime import date
from ftplib import error_perm
from io import BytesIO
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from config.settings import CONFIDENTIALITY_TEMPLATE, INFRASTRUCTURE_TEMPLATE
from services.document_generation.confidentiality_service import (
    ConfidentialityService,
    DocumentGenerationError,
)
from services.document_generation.infrastructure_service import InfrastructureService
from services.signature_storage import SignatureStorage, SignatureStorageError


PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)
SETTINGS = {
    "type": "ftps",
    "host": "host.josekrea.com",
    "port": 21,
    "username": "test-user",
    "password": "test-password",
    "remote_dir": "/firmas",
}


class FakeFTP:
    def __init__(self, **_kwargs) -> None:
        self.calls: list[tuple] = []
        self.directories = {"/"}
        self.files: dict[str, bytes] = {}
        self.encoding = "ascii"

    def connect(self, host: str, port: int, timeout: int) -> None:
        self.calls.append(("connect", host, port, timeout))

    def login(self, username: str, password: str) -> None:
        self.calls.append(("login", username, password))

    def set_pasv(self, enabled: bool) -> None:
        self.calls.append(("set_pasv", enabled))

    def prot_p(self) -> None:
        self.calls.append(("prot_p",))

    def cwd(self, path: str) -> None:
        self.calls.append(("cwd", path))
        if path not in self.directories:
            raise error_perm("550 Directory unavailable")

    def mkd(self, path: str) -> None:
        self.calls.append(("mkd", path))
        self.directories.add(path)

    def storbinary(self, command: str, stream: BytesIO, blocksize: int) -> None:
        self.calls.append(("storbinary", command, blocksize))
        self.files[command.removeprefix("STOR ")] = stream.read()

    def rename(self, source: str, target: str) -> None:
        self.calls.append(("rename", source, target))
        self.files[target] = self.files.pop(source)

    def retrbinary(self, command: str, callback, blocksize: int) -> None:
        self.calls.append(("retrbinary", command, blocksize))
        filename = command.removeprefix("RETR ")
        if filename not in self.files:
            raise error_perm("550 File unavailable")
        data = self.files[filename]
        midpoint = max(1, len(data) // 2)
        callback(data[:midpoint])
        callback(data[midpoint:])

    def delete(self, filename: str) -> None:
        self.calls.append(("delete", filename))
        if filename not in self.files:
            raise error_perm("550 File unavailable")
        del self.files[filename]

    def quit(self) -> None:
        self.calls.append(("quit",))

    def close(self) -> None:
        self.calls.append(("close",))


class SignatureStorageTests(unittest.TestCase):
    def setUp(self) -> None:
        self.ftp = FakeFTP()
        self.storage = SignatureStorage(
            settings=SETTINGS,
            ftp_factory=lambda **_kwargs: self.ftp,
        )

    def test_save_and_load_use_private_explicit_ftps(self) -> None:
        reference = self.storage.save("firma.png", PNG)

        self.assertTrue(reference.startswith("signature:"))
        self.assertNotIn("http", reference)
        self.assertNotIn(SETTINGS["host"], reference)
        self.assertNotIn(SETTINGS["username"], reference)
        filename = reference.removeprefix("signature:")
        self.assertEqual(self.ftp.files[filename], PNG)
        self.assertIn(("prot_p",), self.ftp.calls)
        self.assertIn(("mkd", "/firmas"), self.ftp.calls)

        loaded = self.storage.load(reference)

        self.assertEqual(loaded.getvalue(), PNG)
        self.assertIn(
            ("connect", "host.josekrea.com", 21, self.storage.TIMEOUT_SECONDS),
            self.ftp.calls,
        )
        self.assertIn(
            ("retrbinary", f"RETR {filename}", self.storage.BLOCK_SIZE),
            self.ftp.calls,
        )

    def test_invalid_image_and_remote_reference_are_rejected(self) -> None:
        with self.assertRaises(SignatureStorageError):
            self.storage.save("firma.png", b"not-an-image")
        with self.assertRaises(SignatureStorageError):
            self.storage.save("firma.png", b"\x89PNG\r\n\x1a\n")
        with self.assertRaises(SignatureStorageError):
            self.storage.load("signature:../firma.png")
        with self.assertRaises(SignatureStorageError):
            self.storage.load("signature:archivo\r\nDELE firma.png")
        self.assertFalse(self.ftp.calls)

    def test_ftps_errors_do_not_expose_credentials(self) -> None:
        self.ftp.login = lambda *_args: (_ for _ in ()).throw(
            error_perm("530 Login incorrect")
        )
        with self.assertRaises(SignatureStorageError) as context:
            self.storage.save("firma.png", PNG)
        self.assertNotIn(SETTINGS["password"], str(context.exception))


class SignatureDocumentTests(unittest.TestCase):
    def test_infrastructure_adds_remote_signature_from_memory(self) -> None:
        document = Document()
        paragraph = document.add_paragraph()
        with patch.object(
            SignatureStorage,
            "open",
            return_value=BytesIO(PNG),
        ):
            InfrastructureService._add_signature(
                paragraph,
                "signature:test.png",
            )

        self.assertEqual(paragraph.alignment, WD_ALIGN_PARAGRAPH.LEFT)
        self.assertEqual(len(paragraph._p.xpath(".//w:drawing")), 1)
        extent = paragraph._p.xpath(".//wp:extent")[0]
        self.assertEqual(extent.get("cx"), str(Inches(1.1)))

    def test_confidentiality_adds_remote_signature_from_memory(self) -> None:
        document = Document()
        row = document.add_table(rows=1, cols=1).rows[0]
        with patch.object(
            SignatureStorage,
            "open",
            return_value=BytesIO(PNG),
        ):
            source = ConfidentialityService._resolve_signature(
                "signature:test.png"
            )
        ConfidentialityService()._fill_signature_row(
            row,
            {
                "name": "Talento de prueba",
                "role": "Talento Titular",
                "document": "CC 123",
                "signature_path": source,
            },
        )

        signature_cell = row.cells[0].tables[0].rows[0].cells[1]
        drawings = signature_cell._tc.xpath(".//w:drawing")
        self.assertEqual(len(drawings), 1)
        extent = signature_cell._tc.xpath(".//wp:extent")[0]
        self.assertEqual(extent.get("cx"), str(Inches(1.65)))

    def test_storage_failure_becomes_document_error(self) -> None:
        document = Document()
        paragraph = document.add_paragraph()
        with patch.object(
            SignatureStorage,
            "open",
            side_effect=SignatureStorageError("Firma no disponible"),
        ):
            with self.assertRaises(DocumentGenerationError):
                InfrastructureService._add_signature(
                    paragraph,
                    "signature:test.png",
                )

    def test_infrastructure_template_accepts_remote_signatures(self) -> None:
        project = {"city": "Bogotá", "code": "PR-001", "name": "Proyecto"}
        titular = {
            "name": "Talento",
            "document_type": "CC",
            "document_number": "123",
            "email": "talento@example.com",
            "signature_path": "signature:talento.png",
        }
        expert = {
            "name": "Experto",
            "document_type": "CC",
            "document_number": "456",
            "email": "experto@example.com",
            "signature_path": "signature:experto.png",
        }
        with tempfile.TemporaryDirectory() as temporary_directory:
            document_path = Path(temporary_directory) / "infraestructura.docx"
            shutil.copy2(INFRASTRUCTURE_TEMPLATE, document_path)
            with patch.object(
                SignatureStorage,
                "open",
                side_effect=lambda _reference: BytesIO(PNG),
            ):
                InfrastructureService()._complete_template(
                    document_path,
                    project,
                    titular,
                    expert,
                    date(2026, 8, 29),
                    date(2026, 1, 15),
                )

            completed = Document(document_path)
            self.assertEqual(
                len(completed.paragraphs[31]._p.xpath(".//w:drawing")),
                1,
            )
            self.assertEqual(
                len(completed.paragraphs[37]._p.xpath(".//w:drawing")),
                1,
            )

    def test_confidentiality_template_accepts_remote_signature(self) -> None:
        project = {"city": "Bogotá", "code": "PR-001", "name": "Proyecto"}
        signers = [
            {
                "person": {
                    "name": "Talento",
                    "document_type": "CC",
                    "document_number": "123",
                    "document_issue_place": "Bogotá",
                    "signature_path": "signature:talento.png",
                },
                "roles": ["titular", "interlocutor", "ejecutor"],
            }
        ]
        with tempfile.TemporaryDirectory() as temporary_directory:
            document_path = Path(temporary_directory) / "confidencialidad.docx"
            shutil.copy2(CONFIDENTIALITY_TEMPLATE, document_path)
            with patch.object(
                SignatureStorage,
                "open",
                side_effect=lambda _reference: BytesIO(PNG),
            ):
                ConfidentialityService()._complete_template(
                    document_path,
                    project,
                    date(2026, 8, 29),
                    signers,
                )

            completed = Document(document_path)
            self.assertEqual(len(completed.tables), 1)
            self.assertGreaterEqual(
                len(completed.tables[0]._tbl.xpath(".//w:drawing")),
                1,
            )


if __name__ == "__main__":
    unittest.main()
